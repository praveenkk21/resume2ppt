import io
from app.models import ResumeData


def export_as_docx(resume: ResumeData) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Remove default margins — set to 1 inch all around
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    def add_name(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text or "")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        return p

    def add_contact_line(resume):
        parts = [
            resume.email,
            resume.phone,
            resume.linkedin,
            resume.location,
        ]
        line = "  ·  ".join(p for p in parts if p)
        if not line:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def add_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_heading(title):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
        # Bottom border on the heading paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def add_body(text, italic=False, size=10.5):
        p = doc.add_paragraph()
        run = p.add_run(text or "")
        run.font.size = Pt(size)
        run.italic = italic
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text or "")
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    # --- Name ---
    add_name(resume.name or "")
    add_contact_line(resume)
    add_hr()

    # --- Summary ---
    if resume.summary:
        add_section_heading("Summary")
        add_body(resume.summary)

    # --- Experience ---
    if resume.experience:
        add_section_heading("Experience")
        for exp in resume.experience:
            title_company = " — ".join(filter(None, [exp.title, exp.company]))
            if title_company:
                p = doc.add_paragraph()
                run = p.add_run(title_company)
                run.bold = True
                run.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(1)
            if exp.dates:
                add_body(exp.dates, italic=True, size=10)
            for bullet in exp.bullets:
                add_bullet(bullet)

    # --- Education ---
    if resume.education:
        add_section_heading("Education")
        for edu in resume.education:
            parts = [edu.degree, edu.school]
            line = " — ".join(p for p in parts if p)
            if line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10.5)
                p.paragraph_format.space_after = Pt(1)
            if edu.dates:
                add_body(edu.dates, italic=True, size=10)
            if edu.gpa:
                add_body(f"GPA: {edu.gpa}", size=10)

    # --- Skills ---
    if resume.skills:
        add_section_heading("Skills")
        add_body(", ".join(resume.skills))

    # --- Certifications ---
    if resume.certifications:
        add_section_heading("Certifications")
        add_body(", ".join(resume.certifications))

    # --- Languages ---
    if resume.languages:
        add_section_heading("Languages")
        add_body(", ".join(resume.languages))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_as_pdf(resume: ResumeData) -> bytes:
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    INDIGO = colors.HexColor("#4f46e5")
    DARK = colors.HexColor("#1a1a2e")
    GRAY = colors.HexColor("#555555")

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "NameStyle",
        parent=styles["Normal"],
        fontSize=22,
        fontName="Helvetica-Bold",
        textColor=DARK,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "ContactStyle",
        parent=styles["Normal"],
        fontSize=9,
        textColor=GRAY,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    section_style = ParagraphStyle(
        "SectionStyle",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica-Bold",
        textColor=INDIGO,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "BodyStyle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=DARK,
        spaceAfter=2,
    )
    bold_body_style = ParagraphStyle(
        "BoldBody",
        parent=body_style,
        fontName="Helvetica-Bold",
        spaceAfter=1,
    )
    italic_style = ParagraphStyle(
        "ItalicStyle",
        parent=body_style,
        fontName="Helvetica-Oblique",
        fontSize=9.5,
        textColor=GRAY,
    )
    bullet_style = ParagraphStyle(
        "BulletStyle",
        parent=body_style,
        leftIndent=14,
        bulletIndent=4,
        spaceAfter=1,
    )

    story = []

    # Name
    story.append(Paragraph(resume.name or "", name_style))

    # Contact line
    contact_parts = [resume.email, resume.phone, resume.linkedin, resume.location]
    contact_line = "  ·  ".join(p for p in contact_parts if p)
    if contact_line:
        story.append(Paragraph(contact_line, contact_style))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CCCCCC"), spaceAfter=6))

    def section(title):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#AAAAAA"), spaceAfter=4))

    # Summary
    if resume.summary:
        section("Summary")
        story.append(Paragraph(resume.summary, body_style))

    # Experience
    if resume.experience:
        section("Experience")
        for exp in resume.experience:
            title_co = " — ".join(filter(None, [exp.title, exp.company]))
            if title_co:
                story.append(Paragraph(title_co, bold_body_style))
            if exp.dates:
                story.append(Paragraph(exp.dates, italic_style))
            for bullet in exp.bullets:
                story.append(Paragraph(f"• {bullet}", bullet_style))
        story.append(Spacer(1, 4))

    # Education
    if resume.education:
        section("Education")
        for edu in resume.education:
            parts = [edu.degree, edu.school]
            line = " — ".join(p for p in parts if p)
            if line:
                story.append(Paragraph(line, bold_body_style))
            if edu.dates:
                story.append(Paragraph(edu.dates, italic_style))
            if edu.gpa:
                story.append(Paragraph(f"GPA: {edu.gpa}", body_style))
        story.append(Spacer(1, 4))

    # Skills
    if resume.skills:
        section("Skills")
        story.append(Paragraph(", ".join(resume.skills), body_style))

    # Certifications
    if resume.certifications:
        section("Certifications")
        story.append(Paragraph(", ".join(resume.certifications), body_style))

    # Languages
    if resume.languages:
        section("Languages")
        story.append(Paragraph(", ".join(resume.languages), body_style))

    doc.build(story)
    return buf.getvalue()
