import io
from docx import Document
from docx.shared import Pt


def extract_paragraph_inventory(file_bytes: bytes) -> list:
    """
    Returns a list of paragraph descriptors from the DOCX template.
    Each entry: {para_index, style_name, current_text, is_in_table, table_index, cell_coords}
    """
    doc = Document(io.BytesIO(file_bytes))
    inventory = []

    # Body paragraphs
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        inventory.append({
            "para_index": idx,
            "source": "body",
            "style_name": para.style.name,
            "current_text": text,
            "is_in_table": False,
        })

    # Table cells
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    inventory.append({
                        "para_index": p_idx,
                        "source": "table",
                        "table_index": t_idx,
                        "row": r_idx,
                        "col": c_idx,
                        "style_name": para.style.name,
                        "current_text": text,
                        "is_in_table": True,
                    })

    return [item for item in inventory if item["current_text"]]


def _replace_para_text(para, new_text: str):
    """
    Replace text in a paragraph while preserving the first run's formatting.
    Uses the same run-safe approach as ppt_handler — never clears runs entirely,
    never sets paragraph.text directly.
    """
    if not para.runs:
        # No runs at all — just set via XML text node if present
        for node in para._p:
            if node.text is not None:
                node.text = new_text
        return

    # Write new text into first run, blank the rest
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def apply_docx_mapping(file_bytes: bytes, mapping: list) -> bytes:
    """
    mapping: list of {source, para_index, table_index?, row?, col?, new_text}
    """
    doc = Document(io.BytesIO(file_bytes))

    for entry in mapping:
        new_text = entry.get("new_text", "")
        if not new_text:
            continue

        if not entry.get("is_in_table", False):
            idx = entry["para_index"]
            if idx < len(doc.paragraphs):
                _replace_para_text(doc.paragraphs[idx], new_text)
        else:
            t_idx = entry.get("table_index", 0)
            r_idx = entry.get("row", 0)
            c_idx = entry.get("col", 0)
            p_idx = entry.get("para_index", 0)
            try:
                cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
                if p_idx < len(cell.paragraphs):
                    _replace_para_text(cell.paragraphs[p_idx], new_text)
            except (IndexError, KeyError):
                continue

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
